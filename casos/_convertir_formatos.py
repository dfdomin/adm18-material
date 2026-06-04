"""Batch convert case .md files to native formats (.eml, .docx, .xlsx, .pdf)."""
import os, re, sys
from datetime import datetime

BASE = os.path.dirname(os.path.abspath(__file__))

# ── .eml generator ──
def gen_eml(md_path, eml_path):
    with open(md_path) as f:
        content = f.read()
    # Extract subject, date, from, to from markdown
    subject = re.search(r'\*\*ASUNTO:?\*\*\s*(.+?)(?:\n|$)', content)
    subject = subject.group(1).strip() if subject else "Sin asunto"
    date_match = re.search(r'\*\*FECHA:?\*\*\s*(.+?)(?:\n|$)', content)
    date = date_match.group(1).strip() if date_match else datetime.now().strftime("%d de %B de %Y")
    to_match = re.search(r'\*\*PARA:?\*\*\s*(.+?)(?:\n|$)', content)
    to_addr = to_match.group(1).strip() if to_match else "cliente@email.com"
    from_match = re.search(r'\*\*De:?\*\*\s*(.+?)(?:\n|$)', content)
    if from_match:
        from_addr = from_match.group(1).strip()
    else:
        from_addr = "servicio@empresa.com"
    
    # Get body: everything after the first --- separator
    parts = content.split('---', 1)
    body = parts[1] if len(parts) > 1 else content
    # Clean up markdown
    body = body.replace('**', '').strip()
    
    eml = f"""From: {from_addr}
To: {to_addr}
Date: {date}
Subject: {subject}
MIME-Version: 1.0
Content-Type: text/plain; charset=utf-8
Content-Transfer-Encoding: 8bit

{body}
"""
    with open(eml_path, 'w', encoding='utf-8') as f:
        f.write(eml)
    return os.path.getsize(eml_path)


# ── .docx generator ──
def gen_docx(md_path, docx_path):
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    with open(md_path) as f:
        content = f.read()
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    lines = content.split('\n')
    for line in lines:
        s = line.strip()
        if not s: continue
        if s.startswith('# ') or s.startswith('## ') or s.startswith('### '):
            heading = doc.add_heading(s.lstrip('#').strip(), level=min(s.count('#'), 3))
        elif s == '---':
            doc.add_paragraph('_' * 60)
        elif s.startswith('|') and '---' in s:
            continue  # skip table separators
        elif s.startswith('|'):
            # Table row
            cells = [c.strip() for c in s.split('|')[1:-1]]
            if cells:
                p = doc.add_paragraph()
                p.style.font.size = Pt(9)
                run = p.add_run('  |  '.join(cells))
                run.font.size = Pt(9)
        elif s.startswith('- '):
            doc.add_paragraph(s[2:], style='List Bullet')
        elif s.startswith('**') and ':**' in s:
            doc.add_paragraph(s.replace('**', '').replace(':*', ':'), style='List Bullet')
        else:
            # Clean formatting
            txt = s.replace('**', '').replace('*', '').replace('`', '')
            if txt:
                doc.add_paragraph(txt)
    
    doc.save(docx_path)
    return os.path.getsize(docx_path)


# ── .xlsx generator ──
def gen_xlsx(md_path, xlsx_path):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    with open(md_path) as f:
        content = f.read()
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Documento"
    
    header_fill = PatternFill(start_color="1a2744", end_color="1a2744", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    body_font = Font(name="Calibri", size=9)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    row_num = 1
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        
        if s.startswith('# '):
            ws.cell(row=row_num, column=1, value=s[2:]).font = Font(name="Calibri", bold=True, size=14, color="1a2744")
            row_num += 2
            i += 1
        elif s.startswith('## '):
            ws.cell(row=row_num, column=1, value=s[3:]).font = Font(name="Calibri", bold=True, size=12, color="2c5ea8")
            row_num += 1
            i += 1
        elif s.startswith('**') and ':**' in s:
            txt = s.replace('**', '').replace(':*', ':')
            ws.cell(row=row_num, column=1, value=txt).font = Font(name="Calibri", bold=True, size=9)
            row_num += 1
            i += 1
        elif s == '---':
            row_num += 1
            i += 1
        elif s.startswith('|') and i+1 < len(lines) and '---' in lines[i+1]:
            # Table
            tl = []; j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                tl.append(lines[j].strip()); j += 1
            
            if len(tl) >= 2:
                hdrs = [c.strip() for c in tl[0].split('|')[1:-1]]
                for ci, h in enumerate(hdrs, 1):
                    cell = ws.cell(row=row_num, column=ci, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.border = thin_border
                    cell.alignment = Alignment(wrap_text=True)
                row_num += 1
                
                for tr in tl[2:]:
                    cells = [c.strip() for c in tr.split('|')[1:-1]]
                    for ci, cv in enumerate(cells, 1):
                        cell = ws.cell(row=row_num, column=ci, value=cv)
                        cell.font = body_font
                        cell.border = thin_border
                        cell.alignment = Alignment(wrap_text=True)
                    row_num += 1
                row_num += 1
            i = j
        elif s and not s.startswith('#'):
            ws.cell(row=row_num, column=1, value=s.replace('**', '')).font = body_font
            row_num += 1
            i += 1
        else:
            i += 1
    
    # Adjust column widths
    for col in ws.columns:
        max_len = 0
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
    
    wb.save(xlsx_path)
    return os.path.getsize(xlsx_path)


# ── Mapping: which docs get which formats ──
CONVERSIONS = {
    # LatamBox
    "latambox/01_correo_bienvenida.md": ["eml"],
    "latambox/02_formulario_registro.md": ["docx"],
    "latambox/03_factura_amazon.md": ["xlsx"],
    "latambox/04_notificacion_llegada.md": ["eml"],
    "latambox/05_declaracion_valor.md": ["docx"],
    "latambox/06_guia_aerea_awb.md": ["xlsx"],
    "latambox/07_factura_flete.md": ["xlsx"],
    "latambox/08_correo_nacionalizacion.md": ["eml"],
    "latambox/09_comprobante_impuestos.md": ["xlsx"],
    "latambox/10_guia_entrega_local.md": ["docx"],
    "latambox/11_correo_confirmacion_entrega.md": ["eml"],
    "latambox/12_encuesta_y_queja.md": ["docx"],
    # AndesBox
    "andesbox/01_reporte_operaciones_julio.md": ["xlsx"],
    "andesbox/02_queja_cliente_gonzalez.md": ["eml"],
    "andesbox/03_queja_cliente_fuentes.md": ["eml"],
    "andesbox/04_factura_flete_lima.md": ["xlsx"],
    "andesbox/05_acta_reunion_julio.md": ["docx"],
    # Envíos del Pacífico
    "envios_pacifico/01_correo_bienvenida.md": ["eml"],
    "envios_pacifico/02_factura_alibaba.md": ["xlsx"],
    "envios_pacifico/03_notificacion_recepcion.md": ["eml"],
    "envios_pacifico/04_factura_flete.md": ["xlsx"],
    "envios_pacifico/05_circular_reempaque.md": ["docx"],
    "envios_pacifico/06_queja_dano.md": ["eml"],
    # GlobalBox
    "globalbox/01_correo_bienvenida.md": ["eml"],
    "globalbox/02_reporte_operaciones.md": ["xlsx"],
    "globalbox/03_queja_panama.md": ["eml"],
    "globalbox/04_queja_mexico.md": ["eml"],
    "globalbox/05_factura_flete_madrid.md": ["xlsx"],
    "globalbox/06_acta_junta.md": ["docx"],
    "globalbox/07_cotizacion_doxis.md": ["xlsx"],
    "globalbox/08_memo_protocolo.md": ["docx"],
}

if __name__ == '__main__':
    stats = {"eml": 0, "docx": 0, "xlsx": 0}
    total = 0
    
    for rel_path, formats in CONVERSIONS.items():
        md_path = os.path.join(BASE, rel_path)
        base_name = os.path.splitext(os.path.basename(rel_path))[0]
        folder = os.path.dirname(os.path.join(BASE, rel_path))
        
        for fmt in formats:
            out_path = os.path.join(folder, f"{base_name}.{fmt}")
            try:
                if fmt == "eml":
                    sz = gen_eml(md_path, out_path)
                elif fmt == "docx":
                    sz = gen_docx(md_path, out_path)
                elif fmt == "xlsx":
                    sz = gen_xlsx(md_path, out_path)
                stats[fmt] += 1
                total += 1
                print(f"  {base_name}.{fmt}: {sz:,} bytes")
            except Exception as e:
                print(f"  {base_name}.{fmt}: ERROR - {e}")
    
    print(f"\nGenerados: {total} archivos ({stats['eml']} .eml + {stats['docx']} .docx + {stats['xlsx']} .xlsx)")
