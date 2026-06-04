from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "ADM18_Planeador_Curricular_Procesamiento_Informacion_2026.docx"


BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DDEFE3"
GOLD = "F7E7B7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
            if row_i == 0 and header:
                set_cell_shading(cell, LIGHT_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE if level <= 2 else "434343")
    return p


def add_note(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    p.add_run("\n" + body)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(4)
    return table


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Planeador curricular ADM18\nProcesamiento de la Informacion en la Organizacion")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Modelo de formacion basado en competencias | 13 semanas de desarrollo + semana 14 de evaluacion final").italic = True

    meta = doc.add_table(rows=4, cols=2)
    meta.cell(0, 0).text = "Modulo"
    meta.cell(0, 1).text = "ADM18 - Procesamiento de la Informacion en la Organizacion"
    meta.cell(1, 0).text = "Programas"
    meta.cell(1, 1).text = "Tecnica Profesional en Operaciones del Comercio Exterior / Operacion de Procesos Empresariales"
    meta.cell(2, 0).text = "Caso integrador"
    meta.cell(2, 1).text = "Casillero Caribe Express: de la solicitud del cliente a la decision administrativa"
    meta.cell(3, 0).text = "Fecha base"
    meta.cell(3, 1).text = "Semana 1: 18 al 23 de mayo de 2026. Parcial 1: 15-20 de junio. Parcial 2: 21-27 de julio."
    style_table(meta)

    add_note(
        doc,
        "Decision pedagogica central",
        "El curso usa un solo caso formativo repetido con complejidad creciente. Paneles solares y vehiculos electricos chinos quedan como ejemplos breves o casos espejo, no como eje, para evitar sobrecarga tecnica y regulatoria.",
        GREEN,
    )

    add_heading(doc, "1. Proposito formativo y enfoque de competencia", 1)
    doc.add_paragraph(
        "Este planeador convierte el microcurriculo en una ruta de evidencias observables. El modulo no se limita a reconocer documentos: entrena al estudiante para procesar informacion administrativa, organizarla, comunicarla y usarla como soporte de decisiones en contextos relacionados con comercio exterior y procesos empresariales."
    )
    add_note(
        doc,
        "RA de programa puente ajustable",
        "Como no se adjunto la rubrica oficial del programa, se propone este RA puente para alinear la medicion: 'Gestiona informacion documental de operaciones empresariales y de comercio exterior, aplicando criterios de organizacion, comunicacion, trazabilidad, uso de TIC y soporte a la toma de decisiones'. Reemplazar este texto por el RA oficial cuando este disponible.",
        GOLD,
    )

    add_heading(doc, "2. Trazabilidad microcurricular", 1)
    trace = doc.add_table(rows=1, cols=6)
    headers = ["RA modulo", "Criterios microcurriculo", "Contenido visible en el plan", "Evidencia", "RA programa puente", "Indicador de rubrica"]
    for i, h in enumerate(headers):
        set_cell_text(trace.cell(0, i), h, True)
    rows = [
        [
            "RA1. Identificar documentos administrativos comunes.",
            "CE1 clasifica; CE2 interpreta partes; CE3 distingue concepto, uso y clases.",
            "Semanas 1-3: ciclo de vida documental, funcion administrativa, tipologias.",
            "Matriz documento-funcion del caso Casillero Caribe Express.",
            "Gestiona informacion documental.",
            "PI1. Clasifica documentos segun funcion, origen, destinatario y proposito.",
        ],
        [
            "RA2. Construir documentos administrativos bajo estandares de calidad.",
            "CE1 elabora croquis; CE2 digita documentos; CE3 usa lenguaje claro y administrativo.",
            "Semanas 4, 6 y 7: GTC 185 como checklist, redaccion, actas, correos, memos, cartas.",
            "Paquete documental corregido y version final.",
            "Comunica informacion administrativa.",
            "PI2. Produce documentos claros, pertinentes, trazables y ajustados a formato.",
        ],
        [
            "RA3. Analizar procesos de documentacion de planeacion estrategica.",
            "CE1 interpreta planes; CE2 identifica pasos de decision; CE3 compila documentacion de comites.",
            "Semanas 8 y 9: quejas, demoras, evidencias, matrices, memo de decision.",
            "Memo de decision con hallazgos y recomendacion.",
            "Soporta decisiones de operacion.",
            "PI3. Extrae informacion relevante y justifica una decision administrativa.",
        ],
        [
            "RA4. Valorar la relevancia de TIC en gestion de informacion.",
            "CE1 describe TIC; CE2 opera conversion digital; CE3 reconoce equipos/recursos.",
            "Semanas 11-13: carpetas, metadatos, formularios, IA asistida, trazabilidad.",
            "Mini sistema documental digital + checklist de portafolio.",
            "Usa TIC con criterio documental.",
            "PI4. Organiza evidencias digitales con trazabilidad, versionamiento y uso etico de IA.",
        ],
    ]
    for row in rows:
        cells = trace.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(trace)

    add_heading(doc, "3. Rubrica puente de programa", 1)
    rub = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Indicador", "Inicial", "Basico", "Competente", "Destacado"]):
        set_cell_text(rub.cell(0, i), h, True)
    rub_rows = [
        ["PI1 Clasificacion documental", "Nombra documentos sin criterio estable.", "Clasifica algunos documentos con ayuda.", "Clasifica por funcion, origen, destinatario y proposito.", "Justifica clasificacion y detecta inconsistencias."],
        ["PI2 Produccion documental", "Redacta con proposito ambiguo o formato incompleto.", "Cumple estructura basica con errores menores.", "Produce documento claro, pertinente y formal.", "Ajusta tono, formato y evidencia al destinatario."],
        ["PI3 Analisis para decision", "Enumera datos sin relacionarlos.", "Identifica hechos relevantes con poca priorizacion.", "Organiza hallazgos y recomienda una accion viable.", "Relaciona riesgos, evidencias, causas y consecuencias."],
        ["PI4 TIC y trazabilidad", "Guarda archivos sin orden verificable.", "Usa nombres o carpetas parcialmente consistentes.", "Organiza documentos con version, fecha y evidencia.", "Propone mejoras digitales y uso responsable de IA."],
    ]
    for row in rub_rows:
        cells = rub.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(rub)

    add_heading(doc, "4. Caso integrador", 1)
    doc.add_paragraph(
        "Casillero Caribe Express es una empresa simulada que recibe compras internacionales de clientes, consolida paquetes, gestiona comunicaciones, soportes, reclamos, reportes internos y decisiones de mejora. El caso permite trabajar comercio exterior desde la operacion documental sin abrir una carga excesiva de aranceles, homologaciones o regulacion tecnica."
    )
    for item in [
        "Familia documental 1: documentos de comunicacion: correos, cartas, memorandos, circulares.",
        "Familia documental 2: documentos de soporte operativo: solicitudes, registros, facturas soporte, reportes de seguimiento.",
        "Familia documental 3: documentos de decision: actas, matrices de hallazgos, informes breves y memo de recomendacion.",
    ]:
        bullet(doc, item)

    add_heading(doc, "5. Cronograma de 13 semanas", 1)
    plan = doc.add_table(rows=1, cols=8)
    for i, h in enumerate(["Sem", "Fecha", "RA", "Tema microcurricular", "Actividad/caso", "Evidencia", "Evaluacion", "No preparar"]):
        set_cell_text(plan.cell(0, i), h, True)
    weeks = [
        ["1", "18-23 may", "RA1", "Procesamiento de informacion, documentos y cerebro", "Diagnostico + Quizizz: como pienso y proceso informacion", "Diagnostico no calificable", "Autoevaluacion", "Neurociencia profunda"],
        ["2", "25-30 may", "RA1", "Ciclo de vida y clasificacion documental", "Recuperacion Teams/asinc: 12 documentos del casillero", "Quiz + matriz", "Autoevaluable", "Clase larga grabada"],
        ["3", "1-6 jun", "RA1", "Clases, usos y componentes de documentos", "Clasificar solicitudes, quejas, actas, correos y soportes", "Matriz documento-funcion", "Checklist", "Normas completas"],
        ["4", "8-13 jun", "RA2", "GTC 185, redaccion y formatos", "Corregir documentos mal redactados", "Carta/correo/memo", "Rubrica corta", "Historia de ICONTEC"],
        ["5", "15-20 jun", "RA1-RA2", "Parcial 1 sin tema nuevo", "Caso espejo: EnviOS del Norte", "Instrumento parcial", "Calificacion", "Contenido nuevo"],
        ["6", "22-27 jun", "RA2", "Calidad documental: claridad, tono y proposito", "Responder reclamos de clientes", "Documento limpio", "Coevaluacion", "Feedback extenso"],
        ["7", "29 jun-4 jul", "RA2", "Actas, comunicaciones internas e informes breves", "Reunion por retrasos y circular interna", "Acta + circular", "Rubrica 4 criterios", "Mas de 2 formatos"],
        ["8", "6-11 jul", "RA3", "Informacion para toma de decisiones", "Extraer hechos de quejas, demoras y costos", "Tabla de hallazgos", "Auto/IA asistida", "Analisis financiero complejo"],
        ["9", "13-18 jul", "RA3", "Memo de decision y plan de mejora", "Priorizar problema operativo", "Memo de decision", "Rubrica corta", "Teoria estrategica amplia"],
        ["10", "20-25 jul", "RA2-RA3", "Parcial 2 sin tema nuevo", "Caso espejo: BoxAndes Express", "Analisis documental + memo", "Calificacion", "Tema nuevo"],
        ["11", "27 jul-1 ago", "RA4", "TIC, carpetas, versionamiento y trazabilidad", "Disenar estructura documental digital", "Mapa de archivo", "Checklist", "Herramientas nuevas de mas"],
        ["12", "3-8 ago", "RA4", "IA asistida y etica documental", "Mejorar redaccion con IA sin inventar hechos", "Documento + bitacora", "Auto/coevaluacion", "Debate teorico largo"],
        ["13", "10-15 ago", "RA1-RA4", "Integracion y simulacro final", "Portafolio minimo + defensa breve", "Checklist portafolio", "Revision en clase", "Finalizar en casa"],
    ]
    for row in weeks:
        cells = plan.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(plan)

    add_heading(doc, "6. Semana 1: cerebro, informacion y diagnostico", 1)
    doc.add_paragraph(
        "La primera clase conecta el nombre del modulo con la experiencia real del estudiante: antes de procesar informacion en una organizacion, conviene reconocer como atendemos, filtramos, codificamos, recordamos y comunicamos informacion. El enfoque es metacognitivo, no clinico."
    )
    for item in [
        "Idea 1: el cerebro no registra todo; selecciona por atencion, emocion, experiencia previa y relevancia.",
        "Idea 2: la memoria de trabajo es limitada; por eso los documentos ordenan lo que la mente no debe cargar sola.",
        "Idea 3: comprender no siempre significa recordar; se necesita recuperacion, ejemplos y uso.",
        "Idea 4: procesar informacion en una empresa exige convertir datos sueltos en documentos utiles para otra persona.",
    ]:
        bullet(doc, item)

    add_heading(doc, "7. Prueba diagnostica y juego Quizizz", 1)
    diag = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Item", "Pregunta", "Tipo", "Uso pedagogico"]):
        set_cell_text(diag.cell(0, i), h, True)
    q_rows = [
        ["1", "Cuando recibes mucha informacion nueva, que haces primero?", "Autopercepcion", "Reconocer estrategia espontanea."],
        ["2", "Que diferencia hay entre dato, informacion y documento?", "Conceptual", "Detectar punto de partida del modulo."],
        ["3", "Si una empresa recibe una queja por entrega tardia, que documento deberia producir?", "Situacional", "Conectar cerebro-contexto-documento."],
        ["4", "Que te ayuda mas a recordar: repetir, explicar, aplicar o relacionar?", "Metacognitiva", "Hablar de aprendizaje sin moralizar."],
        ["5", "Un correo mal escrito puede afectar una decision empresarial. Verdadero o falso?", "V/F", "Abrir RA2 y RA3."],
        ["6", "Que evidencia permite saber que una decision no fue improvisada?", "Situacional", "Abrir trazabilidad."],
        ["7", "Que nombre de archivo facilita recuperar informacion despues?", "Opcion multiple", "Abrir RA4."],
        ["8", "Si IA mejora un texto pero inventa un dato, el documento es confiable?", "V/F", "Abrir uso etico de TIC."],
        ["9", "Que parte del cerebro simboliza mejor la atencion: foco, bodega, autopista o archivo?", "Analogica", "Activacion ludica."],
        ["10", "Que necesitas aprender sobre tu forma de pensar para procesar mejor informacion?", "Respuesta corta", "Cierre reflexivo."],
    ]
    for row in q_rows:
        cells = diag.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(diag)

    add_heading(doc, "8. Flujo de evaluacion para evitar acumulacion", 1)
    for item in [
        "Todo lo que sea clasificacion o seleccion se monta como quiz autoevaluable.",
        "Todo documento de practica se revisa con checklist o coevaluacion antes de llegar al docente.",
        "El docente solo califica un producto corto por corte, con la rubrica puente de 4 indicadores.",
        "La retroalimentacion docente se limita a una fortaleza, una correccion prioritaria y una accion siguiente.",
        "Las semanas 5 y 10 no desarrollan contenido nuevo; son cierre de evidencia y medicion.",
    ]:
        numbered(doc, item)

    add_heading(doc, "9. Evaluaciones espejo", 1)
    evals = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Momento", "Caso", "RA medido", "Tarea", "Por que es espejo"]):
        set_cell_text(evals.cell(0, i), h, True)
    eval_rows = [
        ["Parcial 1", "Envios del Norte", "RA1-RA2", "Clasificar documentos, corregir formato y redactar respuesta.", "Misma logica que Casillero Caribe, datos cambiados."],
        ["Parcial 2", "BoxAndes Express", "RA2-RA3", "Analizar carpeta documental y escribir memo de decision.", "Misma familia documental, nueva situacion."],
        ["Final semana 14", "Casillero internacional de accesorios", "RA1-RA4", "Organizar dossier, justificar decision y proponer trazabilidad digital.", "Integra las mismas competencias sin tema nuevo."],
    ]
    for row in eval_rows:
        cells = evals.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(evals)

    doc.add_page_break()
    add_heading(doc, "Anexo A. Lista corta de preparacion docente semanal", 1)
    for item in [
        "Elegir el documento de la semana.",
        "Preparar una plantilla editable.",
        "Preparar un ejemplo correcto y uno con errores.",
        "Crear o reutilizar un quiz de maximo 10 preguntas.",
        "Definir una sola evidencia y una rubrica de maximo 4 criterios.",
    ]:
        bullet(doc, item)

    add_heading(doc, "Anexo B. Prompts docentes para IA", 1)
    doc.add_paragraph("Prompt para generar caso espejo:")
    doc.add_paragraph(
        "Crea un caso breve de empresa de casillero/logistica con 6 documentos administrativos mezclados. Debe evaluar clasificacion documental, redaccion clara, trazabilidad y una decision administrativa. No incluyas datos tecnicos de aduanas ni calculos complejos.",
        style=None,
    )
    doc.add_paragraph("Prompt para revisar un documento estudiantil:")
    doc.add_paragraph(
        "Evalua este documento con cuatro criterios: proposito, formato, claridad y utilidad para decidir. Devuelve una fortaleza, una correccion prioritaria y una accion siguiente. No escribas comentarios largos.",
        style=None,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
